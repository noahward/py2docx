
how_to_use = """ # - Example run from command line: 'python py2docx.py analysis.py'
                 # - File you're converting must be a .py, and be in your command line current working directory
                 # - Your python file must call "fig.savefig('./yourFigName')" if you want figures in your document
                 # - Script runs and saves python file figures, then transfers everything to a word document
                 # - How to format your .py file:
                        # Page title - line starts with #####
                        # Header - line starts with ###
                        # Subtitle - line starts with ####
                        # Unordered list - line starts with # -
                        # Regular comments - line starts with #
             """

# Path to where your final report will end up
FINAL_REPORT = "./report.docx" 

import os, sys
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor


# Function to format body and titles and insert images
def para(doc, line, line_list):
    # Page title
    if (line.startswith("#####")):
        line = line[6:]
        h = doc.add_heading(line, 0)
        return h
    # Header
    if (line.startswith("###")):
        line = line[4:]
        a = doc.add_heading(line, level=1)
        return a
    # Subheader
    if (line.startswith("####")):
        line = line[5:]
        b = doc.add_heading(line, level=2)
        return b
    # Unordered list
    if (line.startswith("# -")):
        line = line[4:]
        c = doc.add_paragraph(line, style='List Bullet')
        c.paragraph_format.space_before = Pt(0)
        c.paragraph_format.space_after = Pt(0)
        return c
    # Insert figure
    if ("fig.savefig" in line):
        line = line[13:-3]
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(line, width=Inches(8))
        return r
    # Comments and code formatting
    else:
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(0)
        d.paragraph_format.space_after = Pt(0)
        i=0
        while i<len(line_list):
            if (line.startswith("#")):
                run = d.add_run(line)
                font = run.font
                font.color.rgb = RGBColor(0, 102, 0)
                i+=1
            else:
                d.add_run(line)
                i+=1
            return d

        
# Ensure user is entering the correct thing into command line
def check_args(argv):
    _, file = argv
    if (file[:-2] != '.py'):
        print("File must be a python file (.py)")
        return -1
    if (len(argv) != 2):
        print("""In your command line, enter:
                 py2pdf2.py 'yourPyFileNameHere'
              """)
        return -1


def main(argv):
    
    if check_args(argv) == -1:
        return -1
    
    _, filepath = argv
    
    # Execute script, and read each line into its own list
    exec(open(filepath).read())
    file = open(filepath, "r")
    line_list = file.readlines()
    file.close()


    # Create and setup page
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier'
    font.size = Pt(10)

    for line in line_list:  
        para(doc, line, line_list)

    doc.save(FINAL_REPORT)
    
    # Delete figure files
    for files in os.listdir(os.getcwd()):
        if files.endswith(".png") or filename.endswith(".jpg"): 
            os.remove(files)
            continue
        else:
            continue
    
    return 0

if __name__ == "__main__":
    exit(main(sys.argv))